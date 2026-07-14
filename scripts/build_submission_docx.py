#!/usr/bin/env python3
"""Build ICAIS submission Word document with clickable [n] citation links."""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
FIG = ROOT / "figures"
OUT = ROOT / "ICAIS2026_Track2_少年科学家投稿.docx"

CITE_RE = re.compile(r"\[(\d+)\]")


def set_run_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def _style_element(font, size, bold):
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), font)
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    return r_pr


def add_hyperlink_to_bookmark(paragraph, text, bookmark_name, font="宋体", size=12, bold=False):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    r_pr = _style_element(font, size, bold)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(r_pr)
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_para(
    doc,
    text,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=12,
    bold=False,
    font="宋体",
    link_cites=False,
    line_spacing=1.25,
    space_after=0,
    space_before=0,
    first_line_indent=None,
):
    p = doc.add_paragraph()
    p.alignment = align
    if link_cites and CITE_RE.search(text):
        parts = re.split(r"(\[\d+\])", text)
        for part in parts:
            if not part:
                continue
            m = CITE_RE.fullmatch(part)
            if m:
                add_hyperlink_to_bookmark(p, part, f"ref{m.group(1)}", font, size, bold)
            else:
                r = p.add_run(part)
                set_run_font(r, font, size, bold)
    else:
        r = p.add_run(text)
        set_run_font(r, font, size, bold)
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    return p


def add_external_hyperlink(paragraph, text, url, font="宋体", size=10):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = _style_element(font, size, False)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(r_pr)
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_ref_para(doc, ref_num, text, url, size=9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bid = 100 + ref_num
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), f"ref{ref_num}")
    p._p.append(start)

    add_external_hyperlink(p, text, url, size=size)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    p._p.append(end)

    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    return p


def add_section(doc, text):
    add_para(doc, text, bold=True, size=12, space_before=4, space_after=2)


def add_figure(doc, png_name, caption, width_cm=11.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(0)
    p.add_run().add_picture(str(FIG / png_name), width=Cm(width_cm))
    add_para(
        doc,
        caption,
        WD_ALIGN_PARAGRAPH.CENTER,
        9,
        space_after=3,
        link_cites="[" in caption,
    )


def add_algorithm_block(doc, title, lines):
    add_para(doc, title, bold=True, size=10, space_before=2, space_after=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.1
    pf.space_after = Pt(3)
    pf.space_before = Pt(0)
    pf.left_indent = Cm(0.4)
    text = "\n".join(lines)
    r = p.add_run(text)
    set_run_font(r, "Courier New", 8)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, spec in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        if spec:
            el.set(qn("w:val"), spec.get("val", "single"))
            el.set(qn("w:sz"), str(spec.get("sz", 8)))
            el.set(qn("w:color"), spec.get("color", "000000"))
            el.set(qn("w:space"), "0")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tc_pr.append(borders)


def _style_three_line_table(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)

    n_rows = len(table.rows)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            top = {"val": "single", "sz": 12} if ri == 0 else None
            bottom = None
            if ri == 0:
                bottom = {"val": "single", "sz": 8}
            elif ri == n_rows - 1:
                bottom = {"val": "single", "sz": 12}
            _set_cell_border(cell, top=top, bottom=bottom)


def add_table(
    doc,
    caption,
    headers,
    rows,
    *,
    highlight_last=False,
    center_cols=None,
    caption_above=True,
):
    if caption_above:
        add_para(
            doc,
            caption,
            WD_ALIGN_PARAGRAPH.CENTER,
            10,
            True,
            space_before=2,
            space_after=2,
            link_cites=True,
        )
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, "宋体", 9, True)
    center_cols = center_cols or list(range(1, len(headers)))
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        is_hi = highlight_last and ri == len(rows) - 1
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                para.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if ci in center_cols else WD_ALIGN_PARAGRAPH.LEFT
                )
                for r in para.runs:
                    set_run_font(r, "宋体", 9, is_hi)
    _style_three_line_table(table)
    if not caption_above:
        add_para(
            doc,
            caption,
            WD_ALIGN_PARAGRAPH.CENTER,
            9,
            space_after=3,
            link_cites=True,
        )
    add_para(doc, "", size=1, space_after=2)


REFERENCES = [
    (
        1,
        "[1] Hao, S., et al. Training Large Language Models to Reason in a Continuous Latent Space. arXiv:2412.06769, 2024.",
        "https://arxiv.org/abs/2412.06769",
    ),
    (
        2,
        "[2] Wei, J., et al. Chain-of-Thought Prompting Elicits Reasoning in Large Language Models. NeurIPS, 2022. arXiv:2201.11903.",
        "https://arxiv.org/abs/2201.11903",
    ),
    (
        3,
        "[3] Hao, S., et al. ProsQA (Proof with Search Question-Answering). In: Training Large Language Models to Reason in a Continuous Latent Space. arXiv:2412.06769, 2024.",
        "https://arxiv.org/abs/2412.06769",
    ),
    (
        4,
        "[4] Schuster, T., et al. Confident Adaptive Language Modeling. NeurIPS, 2022. arXiv:2207.07061.",
        "https://arxiv.org/abs/2207.07061",
    ),
    (
        5,
        "[5] Zhu, H., et al. Reasoning by Superposition: A Theoretical Perspective on Chain of Continuous Thought. arXiv:2505.12514, 2025.",
        "https://arxiv.org/abs/2505.12514",
    ),
    (
        6,
        "[6] Goyal, S., et al. Thoughts Are All You Need: Exploiting Latent Space for LLM Reasoning. arXiv:2311.01465, 2023.",
        "https://arxiv.org/abs/2311.01465",
    ),
    (
        7,
        "[7] Pfau, J., et al. Let's Think Dot by Dot: Hidden Computation in Transformer Language Models. arXiv:2404.15758, 2024.",
        "https://arxiv.org/abs/2404.15758",
    ),
    (
        8,
        "[8] Zhou, W., et al. Efficient Prompting via Dynamic Early Exiting in Large Language Models. arXiv:2310.07463, 2023.",
        "https://arxiv.org/abs/2310.07463",
    ),
]


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)

    add_para(
        doc,
        "置信度回退：连续潜在推理中的选择性停步方法",
        WD_ALIGN_PARAGRAPH.CENTER,
        13,
        True,
        space_after=2,
    )
    add_para(doc, "庞淞阳", WD_ALIGN_PARAGRAPH.CENTER, 11, True, space_after=0)
    add_para(
        doc,
        "北京少年人工智能学院 / 北京市海淀区教师进修学校附属实验学校",
        WD_ALIGN_PARAGRAPH.CENTER,
        10,
        space_after=0,
    )
    add_para(doc, "指导教师：付靖文", WD_ALIGN_PARAGRAPH.CENTER, 10, space_after=4)

    add_section(doc, "摘要")
    add_para(
        doc,
        "连续潜在推理（如 Coconut）可在隐空间中维护并行搜索前沿，优于离散链式思维；但现有停步策略要么对全员固定 latent 步数，要么对全员启用细粒度邻域搜索，难以在实例异构下准确部署，且推理代价偏高。本文提出可部署停步方法 confidence_fallback：保留结构路由主路径，仅在 M2 停步头不确定时触发 kNN–M2 备用搜索。在 Coconut/ProsQA 上（419 题、91 项受控实验），τ=0.48 时准确率达 95.23%、回退率 7.2%，较 fixed_3 提升 11.4 个百分点；跨 53 个 OOD 切片，tri_zone 与 hybrid 切片路由获加权 +2.08 个百分点。",
        link_cites=True,
        first_line_indent=0.74,
    )
    add_para(
        doc,
        "关键词  置信度回退；连续潜在推理；自适应停步；Coconut；ProsQA",
        size=10,
        space_after=2,
    )

    add_section(doc, "1  引言")
    add_para(
        doc,
        "大语言模型的链式思维（CoT）以离散 token 逐步推理[2]，步数难以按题调节。Coconut 在隐空间执行连续思维推理[1]，其潜向量可编码多个搜索前沿的叠加态，从而支持隐式并行广度优先搜索（BFS）[5]；在 ProsQA 图可达性任务上，连续思维亦显著优于离散 CoT[3]。已有工作亦探索隐空间推理与早停机制[6][7][8]。",
        link_cites=True,
        first_line_indent=0.74,
    )
    add_para(
        doc,
        "然而，多数方法仍关注固定步数预算或训练期机制，对推理时不使用标签的可部署停步策略研究不足[4]。一类方法用结构路由匹配题深，但单一路径难以覆盖低置信样本；另一类对全员启用邻域在线停步，可靠但推理代价高。CALM[4] 以输出置信度调节计算量，但 Coconut 场景还须将结构路由、停步判别与备用搜索组合为完整流程。",
        link_cites=True,
        first_line_indent=0.74,
    )
    add_para(
        doc,
        "本文聚焦资源高效的可部署停步策略：以结构路由提供主路径，以 M2 置信度门控决定是否对少数样本启用备用搜索。前期实验表明最优连续思维步数与 BFS 推理深度显著相关（r≈0.543）；fixed_3 仅 83.8%，结构路由可提升至 93.6%，超过题深继续加步会 overthink（3 步 83.8% vs 5 步 74.0%）。基于上述观察，本文提出 confidence_fallback，并在同源与跨分布场景下验证其优势。",
        link_cites=True,
        first_line_indent=0.74,
    )
    add_para(
        doc,
        "本文贡献如下：（1）提出 confidence_fallback，以 M2 置信度门控主/备双路径；（2）在 ProsQA 上达 95.23%、回退率 7.2%，优于全部对照；（3）引入 tri_zone 与 hybrid_slice_router，在 53 个 OOD 切片上获加权 +2.08 个百分点。",
        first_line_indent=0.74,
    )

    add_section(doc, "2  方法")
    add_para(doc, "2.1  问题设定", size=10.5, bold=True, space_after=2)
    add_para(
        doc,
        "考虑 ProsQA 图推理任务[3]：给定有向图与根节点，判断两候选目标中哪一个可达。每题需确定连续思维步数 n（latent steps）并输出答案。部署约束为：（1）不预先扫 1–8 步逐一试探；（2）平均每题 latent 步数 mean_n≤4.5。实验基于 Coconut checkpoint_300——即 Coconut 连续思维模型在 ProsQA 上微调至第 300 步保存的权重快照，全文推理固定使用该模型，在 419 题全量集及 53 个 OOD 切片上评估。",
        link_cites=True,
        first_line_indent=0.74,
    )
    add_para(doc, "2.2  confidence_fallback", size=10.5, bold=True, space_before=2, space_after=2)
    add_para(
        doc,
        "confidence_fallback 与 ProsQA 推理环境交互，包含主路径与门控备用路径（图 1）。主路径：由 BFS 估计题深 d，令 n₀=clamp(d)，单次前向得 pred₀。置信度：M2 停步判别头输出 prob₀=sigmoid(M2(h_{n₀},n₀,x))。门控：若 prob₀≥τ 则输出 pred₀；否则触发 kNN+M2 备用路径。阈值 τ=0.48 由验证集 0.42–0.55 扫参确定。",
        first_line_indent=0.74,
    )
    add_figure(
        doc,
        "figure1_confidence_fallback_flow.png",
        "图 1  confidence_fallback 与 ProsQA 推理环境交互。输入 x 含图与可达性查询；主路径得 ŷ₀ 与 p₀；p₀≥τ 则输出 ŷ₀，否则启用 kNN–M2 回退。",
        width_cm=12.5,
    )
    add_algorithm_block(
        doc,
        "算法 1  confidence_fallback 推理流程",
        [
            "输入：x, τ, Coconut, M2, kNN    输出：y",
            "1: n₀ ← max(min_n, min(BFS_depth(x), cap)); pred₀ ← Coconut.forward(x, n₀)",
            "2: prob₀ ← sigmoid(M2(h_{n₀}, n₀, x))",
            "3: if prob₀ < τ then y ← KNN_M2_online_stop(x) else y ← pred₀",
            "4: return y",
        ],
    )
    add_para(doc, "2.3  三区门控与跨分布部署", size=10.5, bold=True, space_before=2, space_after=2)
    add_para(
        doc,
        "同源 ProsQA 以单一阈值 τ=0.48 定稿。面对分布偏移时，单一 τ 易在部分 OOD 切片上误触发回退。本文引入三区门控 tri_zone（t_low=0.40，t_mid=0.48）：prob₀≥0.48 时采用主路径；prob₀<0.40 时回退；0.40≤prob₀<0.48 时仅在答案翻转条件下回退。进一步与切片规则组合为 hybrid_slice_router，对已知易 hurt 切片采用 skip/agreement 策略，其余切片默认 tri_zone。",
        first_line_indent=0.74,
    )

    add_section(doc, "3  实验")
    add_para(doc, "3.1  实验设置", size=10.5, bold=True, space_after=2)
    add_para(
        doc,
        "在 ProsQA 419 题全量集及 53 个 OOD 切片上评估。对比五种方法：fixed_3、auto_route、structure_d、knn_min3 与 confidence_fallback。评价指标为全量准确率、相对 fixed_3 的增量 Δ，以及回退率。基线分为三类：× 固定预算（fixed_3）；✓ 本文复现对照（auto_route、structure_d、knn_min3）；• 本文方法（confidence_fallback 及 tri_zone/hybrid 变体）。",
        first_line_indent=0.74,
    )
    add_para(doc, "3.2  主要结果", size=10.5, bold=True, space_before=2, space_after=2)
    add_para(
        doc,
        "图 2 与表 1 汇总 ProsQA 419 题上的主要结果。confidence_fallback 准确率达 95.23%，为五种方法最高；回退率 7.2%，92.8% 样本仅一次主路径前向。较最强单路径 structure_d（93.6%）再高出 1.6 个百分点，较 fixed_3 提升 11.4 个百分点。最大增益出现在混合跳数与低置信子集。",
        first_line_indent=0.74,
    )
    add_figure(
        doc,
        "figure2_main_results_bar.png",
        "图 2  ProsQA 对比（419 题）。(a) 五方法准确率，虚线为 fixed_3；(b) 主路径与回退路径占比。",
        width_cm=10.5,
    )
    add_table(
        doc,
        "表 1  主要方法与基线对比（ProsQA 419 题）。加粗与下划线分别表示最佳与次佳。",
        ["方法", "类型", "准确率(%)", "Δ(pp)", "回退率(%)", "核心机制"],
        [
            ["fixed_3", "× 固定预算", "83.8", "—", "—", "全员相同步数"],
            ["auto_route", "✓ 对照", "93.1", "+9.3", "—", "BFS 按题配步"],
            ["structure_d", "✓ 对照", "93.6", "+9.8", "—", "题深预算一次推理"],
            ["knn_min3", "✓ 对照", "92.6", "+8.8", "—", "kNN 辅助停步"],
            ["confidence_fallback", "• 本文", "95.23", "+11.4", "7.2", "主路径+门控回退"],
        ],
        highlight_last=True,
        center_cols=[0, 1, 2, 3, 4],
    )
    add_para(doc, "3.3  分析与跨分布结果", size=10.5, bold=True, space_before=2, space_after=2)
    add_para(
        doc,
        "期望准确率可分解为 p=p₀(1−f)+p₁·f，其中 p₀、p₁ 为主路径与回退路径准确率，f 为回退率。结构路由将 p₀ 提升至 93.6%，confidence_fallback 以 f=7.2% 的选择性回退将 p 推高至 95.23%。fixed_3 忽略题间异构；auto_route 与 structure_d 解决配步但无力补救停步不稳；knn_min3 对全员付出额外计算。confidence_fallback 保留 structure_d 高效主路径，仅在约 7.2% 低置信样本回退。五种子稳健性（μ=93.89%）表明 τ=0.48 并非个例调参；tri_zone 与 hybrid 路由在 53 个切片上获加权 +2.08 个百分点（OOD 子集 +7.44 个百分点）。",
        first_line_indent=0.74,
    )

    add_section(doc, "4  结论")
    add_para(
        doc,
        "本文提出置信度回退 confidence_fallback，将结构路由、M2 置信度门控与选择性 kNN 备用搜索结合为可部署停步流程。在 ProsQA 上达 95.23%、较 fixed_3 提升 11.4 个百分点，且仅 7.2% 样本触发回退。跨分布场景以 tri_zone 与 hybrid 切片路由协同。实验表明，选择性纠错——而非全员固定步数或全员细停——是连续潜在推理高效部署的关键。",
        first_line_indent=0.74,
    )

    add_section(doc, "参考文献")
    for num, text, url in REFERENCES:
        add_ref_para(doc, num, text, url)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Extract text from user-edited Word doc and sync project submission files."""

from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
INBOX = ROOT / "inbox"
DEFAULT_SRC = INBOX / "庞淞阳-ICAIS2026_Track2_少年科学家投稿.docx"
OUT_TXT = ROOT / "ICAIS2026_Track2_少年科学家投稿.txt"
OUT_DOCX = ROOT / "ICAIS2026_Track2_少年科学家投稿.docx"


def para_plain(p) -> str:
    pieces = []
    for child in p._element:
        tag = child.tag.split("}")[-1]
        if tag == "r":
            t = "".join(n.text or "" for n in child.iter() if n.tag.endswith("}t"))
            if t:
                pieces.append(t)
        elif tag == "hyperlink":
            t = "".join(n.text or "" for n in child.iter() if n.tag.endswith("}t"))
            num = t.strip("[]") or "?"
            pieces.append(f"[{num}]")
    return "".join(pieces).strip()


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    blocks: list[str] = []

    for p in doc.paragraphs:
        t = para_plain(p)
        if t:
            blocks.append(t)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append("")
            blocks.extend(rows)

    text = "\n".join(blocks)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + "\n"


def main() -> int:
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SRC
    if not src.exists():
        print(f"未找到文件: {src}")
        print("请把你的 Word 复制到:")
        print(f"  {DEFAULT_SRC}")
        return 1

    text = extract_docx(src)
    OUT_TXT.write_text(text, encoding="utf-8")
    if src.resolve() != OUT_DOCX.resolve():
        shutil.copy2(src, OUT_DOCX)

    print(f"Synced text -> {OUT_TXT}")
    print(f"Synced docx -> {OUT_DOCX}")
    print("--- preview (first 1200 chars) ---")
    print(text[:1200])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
